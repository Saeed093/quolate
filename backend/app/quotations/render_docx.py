"""Render a quotation version to a client-facing DOCX (python-docx).

Deliberately letterhead-free: the top margin is enlarged so the user can print
onto their own pre-printed letterhead stationery. Internal figures (unit cost,
margin, duty) never appear here — they live in the XLSX buildup instead.
"""
from __future__ import annotations

import io


def _fmt(value: object) -> str:
    if value is None or value == "":
        return ""
    try:
        n = float(value)
    except (TypeError, ValueError):
        return str(value)
    if n == int(n):
        return f"{int(n):,}"
    return f"{n:,.2f}"


def render_quotation_docx(data: dict) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document()
    # Reserve blank space at the top for the user's own letterhead.
    doc.sections[0].top_margin = Inches(2.0)

    ccy = data.get("currency", "")

    doc.add_heading("QUOTATION", level=1)

    meta = doc.add_paragraph()
    meta.add_run(f"Quotation No: {data.get('quote_no', '')}\n").bold = True
    if data.get("issue_date"):
        meta.add_run(f"Date: {data['issue_date']}\n")
    if data.get("validity_days"):
        meta.add_run(f"Valid for: {data['validity_days']} days\n")
    if data.get("title"):
        meta.add_run(f"Re: {data['title']}\n")

    lines = data.get("lines", [])
    table = doc.add_table(rows=1, cols=5)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:  # style not present in the default template
        table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(
        ["#", "Description", "Qty", f"Unit Price ({ccy})", f"Total ({ccy})"]
    ):
        hdr[i].text = text
    for ln in lines:
        cells = table.add_row().cells
        cells[0].text = str(ln.get("line_no", ""))
        desc = ln.get("description") or ""
        if ln.get("spec"):
            desc = f"{desc}\n{ln['spec']}"
        cells[1].text = desc
        cells[2].text = _fmt(ln.get("qty"))
        cells[3].text = _fmt(ln.get("unit_price"))
        cells[4].text = _fmt(ln.get("line_total"))

    doc.add_paragraph()
    totals = doc.add_paragraph()
    totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals.add_run(f"Subtotal: {ccy} {_fmt(data.get('subtotal'))}\n")
    if data.get("gst_enabled"):
        totals.add_run(f"GST: {ccy} {_fmt(data.get('tax_total'))}\n")
    grand = totals.add_run(f"Total: {ccy} {_fmt(data.get('grand_total'))}")
    grand.bold = True

    terms = data.get("terms") or {}
    has_terms = data.get("validity_days") or any(
        terms.get(k) for k in ("payment_terms", "delivery", "warranty", "notes")
    )
    if has_terms:
        doc.add_heading("Terms & Conditions", level=2)
        if data.get("validity_days"):
            doc.add_paragraph(
                f"This quotation is valid for {data['validity_days']} days from the date above.",
                style="List Bullet",
            )
        if terms.get("payment_terms"):
            doc.add_paragraph(f"Payment: {terms['payment_terms']}", style="List Bullet")
        if terms.get("delivery"):
            doc.add_paragraph(f"Delivery: {terms['delivery']}", style="List Bullet")
        if terms.get("warranty"):
            doc.add_paragraph(f"Warranty: {terms['warranty']}", style="List Bullet")
        if terms.get("notes"):
            doc.add_paragraph(str(terms["notes"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
